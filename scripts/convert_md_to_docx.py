from __future__ import annotations

import os
import re
import sys
import tempfile
import urllib.request
import base64
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "Bao_Cao_Nang_Cap_Tram_Can.md"
DOCX_PATH = ROOT / "docs" / "Bao_Cao_Nang_Cap_Tram_Can.docx"

def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    
    # Configure headings
    for style_name, size, color in [
        ("Title", 20, RGBColor(31, 78, 121)),
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11, RGBColor(47, 84, 150)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

def add_formatted_text(paragraph, text: str) -> None:
    # Remove markdown links, convert [Text](URL) -> Text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    
    # Parse bold-italic (***), bold (**), italic (*), and inline code (`)
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)

def fetch_mermaid_image(mermaid_code: str) -> bytes | None:
    # Try mermaid.ink first (GET request with base64)
    print(f"Fetching Mermaid diagram from mermaid.ink...")
    try:
        graph_bytes = mermaid_code.encode("utf-8")
        base64_bytes = base64.urlsafe_b64encode(graph_bytes)
        base64_string = base64_bytes.decode("ascii").replace("=", "")
        url = f"https://mermaid.ink/img/{base64_string}"
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36"
        }
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=30) as response:
            return response.read()
    except Exception as e:
        print(f"Error calling mermaid.ink: {e}")
        
    # Try Kroki as fallback (POST request)
    print(f"Fetching Mermaid diagram from Kroki as fallback...")
    url = "https://kroki.io/mermaid/png"
    try:
        headers = {
            "Content-Type": "text/plain",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36"
        }
        req = urllib.request.Request(
            url,
            data=mermaid_code.encode("utf-8"),
            headers=headers
        )
        with urllib.request.urlopen(req, timeout=30) as response:
            return response.read()
    except Exception as e:
        print(f"Error calling kroki.io: {e}")
        
    return None

def create_docx_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    cols_count = len(headers)
    table = document.add_table(rows=1, cols=cols_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        add_formatted_text(p, header_text.strip())
        p.runs[0].bold = True
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                cell = row_cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                add_formatted_text(p, cell_text.strip())
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    if not md_path.exists():
        print(f"Error: {md_path} does not exist.")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_page_layout(doc)
    configure_styles(doc)

    in_code_block = False
    code_block_lines = []
    code_block_lang = ""
    
    in_table = False
    table_headers = []
    table_rows = []

    # Simple block parser
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code Block fences
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                in_code_block = False
                code_content = "\n".join(code_block_lines)
                
                if code_block_lang == "mermaid":
                    # Fetch and insert mermaid image
                    img_data = fetch_mermaid_image(code_content)
                    if img_data:
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            tmp.write(img_data)
                            tmp_path = tmp.name
                        try:
                            # Add paragraph to center image
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.space_after = Pt(12)
                            p.paragraph_format.space_before = Pt(12)
                            p.add_run().add_picture(tmp_path, width=Cm(15))
                        finally:
                            try:
                                os.remove(tmp_path)
                            except Exception:
                                pass
                    else:
                        # Fallback to text block
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run(code_content)
                        run.font.name = "Consolas"
                        run.font.size = Pt(9.5)
                else:
                    # Regular code block
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(code_content)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
                
                code_block_lines = []
                code_block_lang = ""
            else:
                # Start of code block
                in_code_block = True
                code_block_lang = stripped[3:].strip().lower()
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line.rstrip('\r\n'))
            i += 1
            continue

        # Handle Table
        if stripped.startswith("|"):
            # Table row
            in_table = True
            columns = [c.strip() for c in stripped.split("|")[1:-1]]
            
            # Check if it's separator row
            is_separator = False
            if len(columns) > 0 and all(re.match(r'^:?-+:?$', col) for col in columns):
                is_separator = True
                
            if not is_separator:
                if not table_headers:
                    table_headers = columns
                else:
                    table_rows.append(columns)
            i += 1
            continue
        else:
            if in_table:
                # Table ended
                if table_headers:
                    create_docx_table(doc, table_headers, table_rows)
                    # Add spacing after table
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                table_headers = []
                table_rows = []
                in_table = False

        # Handle Headers
        if stripped.startswith("#"):
            # Count hashes
            hash_count = 0
            for char in stripped:
                if char == '#':
                    hash_count += 1
                else:
                    break
            
            header_text = stripped[hash_count:].strip()
            
            if hash_count == 1:
                # Project Title
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.space_before = Pt(12)
                add_formatted_text(p, header_text)
            elif hash_count == 2:
                # Heading 1
                p = doc.add_paragraph(style="Heading 1")
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                add_formatted_text(p, header_text)
            elif hash_count == 3:
                # Heading 2
                p = doc.add_paragraph(style="Heading 2")
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                add_formatted_text(p, header_text)
            elif hash_count == 4:
                # Heading 3
                p = doc.add_paragraph(style="Heading 3")
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                add_formatted_text(p, header_text)
            i += 1
            continue

        # Handle Bullet list items
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            add_formatted_text(p, bullet_text)
            i += 1
            continue

        # Handle horizontal rule
        if stripped == "---":
            # Just ignore horizontal rule or add spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Empty line
        if stripped == "":
            i += 1
            continue

        # Normal Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped)
        i += 1

    # Finalize any remaining table at the end of the file
    if in_table and table_headers:
        create_docx_table(doc, table_headers, table_rows)

    # Save document
    try:
        doc.save(docx_path)
        print(f"Successfully converted {md_path} -> {docx_path}")
    except PermissionError:
        timestamp = datetime.now().strftime("%H%M%S")
        alt_path = docx_path.parent / f"{docx_path.stem}_{timestamp}.docx"
        doc.save(alt_path)
        print(f"Warning: File {docx_path} was locked. Saved copy as: {alt_path}")

if __name__ == "__main__":
    import sys
    md_in = MD_PATH
    docx_out = DOCX_PATH
    if len(sys.argv) > 2:
        md_in = Path(sys.argv[1]).resolve()
        docx_out = Path(sys.argv[2]).resolve()
    elif len(sys.argv) > 1:
        md_in = Path(sys.argv[1]).resolve()
        docx_out = md_in.with_suffix(".docx")
    
    convert_md_to_docx(md_in, docx_out)

