from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "WEIGHING-SESSION-DATA-FLOW.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 20, RGBColor(31, 78, 121)),
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11, RGBColor(47, 84, 150)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title(document: Document) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PHÂN TÍCH LUỒNG DỮ LIỆU WEIGHING SESSION")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Phạm vi: Tạo lượt cân → Cân lần 1 → Cân lần 2 → Phân bổ → Xử lý quá tải → Chuyển xe ra")
    meta.runs[0].italic = True

    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.add_run(f"Ngày xuất tài liệu: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    generated.runs[0].italic = True


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_code_block(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def add_summary_table(document: Document) -> None:
    document.add_heading("1. Bảng Tóm Tắt Luồng", level=1)
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Bước", "Trigger / Hành động", "Bảng bị ghi", "Session status sau bước", "Kết quả nghiệp vụ"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        shade_cell(table.rows[0].cells[index], "D9EAF7")

    rows = [
        [
            "1",
            "Tạo lượt cân",
            "weighing_sessions, weighing_session_lines, vehicle_registrations",
            "PENDING_WEIGHT1",
            "Xe vào phiên cân, đăng ký chuyển sang WEIGHING",
        ],
        [
            "2",
            "Cân lần 1",
            "weighing_sessions, weigh_tickets(master)",
            "PENDING_WEIGHT2",
            "Lưu Weight1 và chốt TTCP 10%",
        ],
        [
            "3",
            "Cân lần 2",
            "weighing_sessions, weigh_tickets(master)",
            "ALLOCATION_PENDING",
            "Tính NetWeight thực tế của lượt cân",
        ],
        [
            "4",
            "Phân bổ thực giao",
            "weighing_session_lines, delivery_tickets, weighing_sessions, weigh_tickets(master)",
            "READY_TO_COMPLETE",
            "Gán thực giao cho từng line và xác định có quá tải hay không",
        ],
        [
            "5",
            "Xử lý quá tải",
            "weighing_sessions, weigh_tickets(split), delivery_tickets(split)",
            "READY_TO_COMPLETE",
            "Chốt trạng thái SPLIT_CONFIRMED hoặc NO_SPLIT_CONFIRMED",
        ],
        [
            "6",
            "Chuyển xe ra",
            "weighing_sessions, vehicle_registrations",
            "COMPLETED",
            "Xe ra bãi, đăng ký sang OUT_YARD",
        ],
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def add_flow_overview(document: Document) -> None:
    document.add_heading("2. Sơ Đồ Luồng Tổng Quan", level=1)
    add_paragraph(
        document,
        "Nguồn chuẩn của workflow là bảng weighing_sessions. Các document như weigh_tickets và delivery_tickets được sinh ra hoặc mirror từ session để phục vụ in ấn, tra cứu và đồng bộ dữ liệu."
    )
    diagram = """
[VehicleRegistration]
REGISTERED + IN_YARD
        |
        v
[Tạo lượt cân]
        |
        v
[WeighingSession]
PENDING_WEIGHT1
        |
        v
[Cân lần 1]
-> cập nhật Weight1
-> tạo/cập nhật Master Weigh Ticket
        |
        v
[WeighingSession]
PENDING_WEIGHT2
        |
        v
[Cân lần 2]
-> cập nhật Weight2, NetWeight
        |
        v
[WeighingSession]
ALLOCATION_PENDING
        |
        v
[Phân bổ thực giao]
-> cập nhật Session Lines
-> tạo/cập nhật Delivery Tickets
-> xác định quá tải
        |
        v
[WeighingSession]
READY_TO_COMPLETE
        |
   +----+----+
   |         |
   | Không   | Có quá tải
   | quá tải | (PENDING)
   |         |
   v         v
[Chuyển xe ra]   [Xử lý quá tải]
                 -> Tách / Không tách
                      |
                      v
                 [Chuyển xe ra]
                      |
                      v
            [VehicleRegistration]
            COMPLETED + OUT_YARD
"""
    add_code_block(document, diagram.strip())


def add_step_detail(
    document: Document,
    title: str,
    purpose: str,
    writes: list[str],
    validations: list[str],
    outputs: list[str],
) -> None:
    document.add_heading(title, level=1)
    add_paragraph(document, purpose, "Mục đích: ")
    document.add_heading("Bảng / dữ liệu bị ghi", level=2)
    add_bullets(document, writes)
    document.add_heading("Điều kiện / kiểm tra chính", level=2)
    add_bullets(document, validations)
    document.add_heading("Trạng thái / kết quả sau bước", level=2)
    add_bullets(document, outputs)


def add_source_of_truth_section(document: Document) -> None:
    document.add_heading("9. Vai Trò Của Từng Bảng Dữ Liệu", level=1)
    items = [
        "vehicle_registrations: nguồn đầu vào nghiệp vụ, theo dõi xe đang ở hàng xe vào, đang cân hay đã ra bãi.",
        "weighing_sessions: nguồn chuẩn cho workflow, chứa trạng thái phiên cân và số cân Weight1 / Weight2 / NetWeight.",
        "weighing_session_lines: snapshot line-level của từng đăng ký trong phiên cân, gồm kế hoạch và thực phân bổ.",
        "weigh_tickets: document phục vụ in ấn, tra cứu, sync; master ticket mirror dữ liệu từ session, split ticket sinh ra khi xử lý quá tải.",
        "delivery_tickets: phiếu thực giao theo line; có loại thường và loại split-derived khi tách quá tải.",
    ]
    add_bullets(document, items)


def add_button_matrix(document: Document) -> None:
    document.add_heading("10. Ma Trận Điều Kiện Nghiệp Vụ Theo Nút UI", level=1)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Nút / hành động", "Điều kiện chính", "Kết quả"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")

    rows = [
        ["Cân lần 1", "SessionStatus = PENDING_WEIGHT1", "Lưu Weight1, tạo/cập nhật master weigh ticket"],
        ["Cân lần 2", "SessionStatus = PENDING_WEIGHT2 và đã có Weight1", "Lưu Weight2, NetWeight"],
        ["Phân bổ", "SessionStatus = ALLOCATION_PENDING hoặc READY_TO_COMPLETE", "Cập nhật allocation line và xác định quá tải"],
        ["Xử lý quá tải", "SessionStatus = READY_TO_COMPLETE, IsOverweight = true, OverweightResolutionStatus = PENDING", "Cho phép tách hoặc xác nhận không tách"],
        ["Chuyển xe ra", "READY_TO_COMPLETE và overweight đã ở trạng thái hợp lệ", "Session COMPLETED, registration OUT_YARD"],
        ["Không lấy hàng", "Phiên chưa COMPLETED/CANCELLED", "Đóng session với NetWeight = 0, hủy documents cũ nếu có"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def build_document() -> Document:
    document = Document()
    set_page_layout(document)
    configure_styles(document)
    add_title(document)

    add_summary_table(document)
    add_flow_overview(document)

    add_step_detail(
        document,
        "3. Bước 1 - Tạo Lượt Cân",
        "Tạo một weighing session mới từ một hoặc nhiều vehicle_registration đang ở hàng xe vào và chuyển các đăng ký đó sang trạng thái đang cân.",
        [
            "Insert weighing_sessions với SessionStatus = PENDING_WEIGHT1.",
            "Insert weighing_session_lines, mỗi registration tạo thành một line với PlannedWeight / PlannedBagCount.",
            "Update vehicle_registrations: RegistrationStatus = IN_SESSION, ProcessingStage = WEIGHING, gắn WeighingSessionId.",
        ],
        [
            "Phải chọn ít nhất một registration.",
            "Registration phải còn tồn tại, chưa hủy, đang ở REGISTERED + IN_YARD.",
            "Không được trộn registration nhập và xuất trong cùng một lượt cân.",
            "Phiếu nhập hiện tại chưa hỗ trợ gộp nhiều registration vào một session.",
        ],
        [
            "Session được tạo và trở thành nguồn chuẩn của workflow.",
            "Chưa phát sinh số cân và chưa bắt buộc có weigh ticket.",
        ],
    )

    add_step_detail(
        document,
        "4. Bước 2 - Cân Lần 1",
        "Ghi nhận trọng lượng lần 1 cho session và chốt snapshot ngưỡng TTCP 10% dùng về sau.",
        [
            "Update weighing_sessions: Weight1, Weight1Time, Ttcp10WeightSnapshot, SessionStatus = PENDING_WEIGHT2.",
            "Insert hoặc update weigh_tickets master với RecordRole = MASTER_SESSION.",
            "Master weigh ticket được mirror từ session qua WeighingSessionTicketSyncService.",
        ],
        [
            "Session phải đang ở PENDING_WEIGHT1.",
            "Nếu cân tay thì tài khoản phải có quyền manual weighing.",
            "TTCP 10% được lấy từ xe hoặc fallback từ tổng PlannedWeight của session lines.",
        ],
        [
            "Session có Weight1 và chờ cân lần 2.",
            "Master weigh ticket tồn tại sau bước này để phục vụ in/tra cứu/sync.",
        ],
    )

    add_step_detail(
        document,
        "5. Bước 3 - Cân Lần 2",
        "Ghi nhận trọng lượng lần 2 và tính NetWeight thực tế của toàn bộ lượt cân.",
        [
            "Update weighing_sessions: Weight2, Weight2Time, NetWeight, SessionStatus = ALLOCATION_PENDING.",
            "Update weigh_tickets master bằng dữ liệu mirror từ session.",
        ],
        [
            "Session phải đang ở PENDING_WEIGHT2 và đã có Weight1.",
            "Phiếu nhập bị chặn nếu Weight2 lớn hơn Weight1.",
            "Nếu cân tay thì tài khoản phải có quyền manual weighing.",
        ],
        [
            "Session có đủ số cân thực tế để bước sang phân bổ.",
            "Chưa xác định quá tải ở bước này; việc đó xảy ra sau phân bổ.",
        ],
    )

    add_step_detail(
        document,
        "6. Bước 4 - Phân Bổ Thực Giao",
        "Phân bổ khối lượng thực giao cho từng line và tạo các delivery ticket thường.",
        [
            "Update weighing_session_lines: ActualAllocatedWeight, ActualAllocatedBagCount, LineStatus = ALLOCATED.",
            "Insert hoặc update delivery_tickets thường với RecordRole = Normal.",
            "Update weighing_sessions: IsOverweight, OverweightAmount, OverweightResolutionStatus, SessionStatus = READY_TO_COMPLETE.",
            "Update weigh_tickets master để mirror NetWeight / Ttcp10WeightSnapshot / IsOverWeight từ session.",
        ],
        [
            "Session phải đang ở ALLOCATION_PENDING hoặc READY_TO_COMPLETE.",
            "Mọi line của session phải có dữ liệu phân bổ.",
            "Tổng ActualAllocatedWeight phải đúng bằng session.NetWeight.",
        ],
        [
            "Nếu NetWeight vượt Ttcp10WeightSnapshot thì session chuyển sang quá tải PENDING.",
            "Nếu không vượt ngưỡng thì session ở NOT_APPLICABLE và sẵn sàng chuyển xe ra.",
        ],
    )

    add_step_detail(
        document,
        "7. Bước 5 - Xử Lý Quá Tải",
        "Xử lý nghiệp vụ khi session đã phân bổ xong nhưng NetWeight vẫn vượt ngưỡng TTCP 10%.",
        [
            "Preview: đọc weighing_sessions, weighing_session_lines và config OverweightSplitStepWeight để dựng phương án split.",
            "Nếu tách: insert weigh_tickets split với RecordRole = SplitDerived.",
            "Nếu tách: insert delivery_tickets split với RecordRole = SplitDerived.",
            "Nếu xử lý lại: soft delete split documents cũ trước khi tạo lại.",
            "Update weighing_sessions: OverweightResolutionStatus = SPLIT_CONFIRMED hoặc NO_SPLIT_CONFIRMED.",
        ],
        [
            "Session phải ở READY_TO_COMPLETE.",
            "Session phải đang IsOverweight = true và OverweightResolutionStatus = PENDING.",
            "Phương án tách phải tạo ra các phiếu con không vi phạm ngưỡng theo rule cấu hình.",
        ],
        [
            "Nhánh Tách: session vẫn là session gốc, split tickets chỉ là document phát sinh.",
            "Nhánh Không tách: không tạo thêm document, chỉ chốt trạng thái xác nhận.",
        ],
    )

    add_step_detail(
        document,
        "8. Bước 6 - Chuyển Xe Ra",
        "Kết thúc phiên cân và đưa toàn bộ registration của session sang màn Danh sách xe ra.",
        [
            "Update weighing_sessions: SessionStatus = COMPLETED.",
            "Update vehicle_registrations: RegistrationStatus = COMPLETED, ProcessingStage = OUT_YARD.",
        ],
        [
            "Session phải READY_TO_COMPLETE.",
            "Session phải có đủ Weight1, Weight2, NetWeight.",
            "OverweightResolutionStatus phải là NOT_APPLICABLE, SPLIT_CONFIRMED hoặc NO_SPLIT_CONFIRMED.",
            "Mọi session line phải ở LineStatus = ALLOCATED và có ActualAllocatedWeight.",
        ],
        [
            "Xe xuất hiện ở màn Danh sách xe ra.",
            "Session kết thúc vòng đời workflow.",
        ],
    )

    add_source_of_truth_section(document)
    add_button_matrix(document)
    return document


def main() -> None:
    document = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
