from __future__ import annotations

import sys
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "Bao_Cao_Nang_Cap_Tram_Can_113702.docx"
MD_OUTPUT_PATH = ROOT / "docs" / "Bao_Cao_Nang_Cap_Tram_Can_from_word.md"

def iter_block_items(parent):
    if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
        
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_markdown_text(paragraph: Paragraph) -> str:
    runs = paragraph.runs
    if not runs:
        return ""
        
    grouped_runs = []
    current_text = ""
    current_bold = None
    current_italic = None
    current_code = None
    
    for run in runs:
        run_text = run.text
        if not run_text:
            continue
        is_bold = bool(run.bold)
        is_italic = bool(run.italic)
        is_code = run.font.name == 'Consolas'
        
        if (current_bold == is_bold and 
            current_italic == is_italic and 
            current_code == is_code):
            current_text += run_text
        else:
            if current_text:
                grouped_runs.append((current_text, current_bold, current_italic, current_code))
            current_text = run_text
            current_bold = is_bold
            current_italic = is_italic
            current_code = is_code
            
    if current_text:
        grouped_runs.append((current_text, current_bold, current_italic, current_code))
        
    # Format grouped runs
    text = ""
    for r_text, bold, italic, code in grouped_runs:
        if bold and italic:
            text += f"***{r_text}***"
        elif bold:
            text += f"**{r_text}**"
        elif italic:
            text += f"*{r_text}*"
        elif code:
            text += f"`{r_text}`"
        else:
            text += r_text
    return text

def format_table_as_md(table: Table) -> str:
    md = []
    if not table.rows:
        return ""
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_texts = []
    for cell in hdr_cells:
        cell_text = " ".join(get_markdown_text(p) for p in cell.paragraphs).strip()
        hdr_texts.append(cell_text.replace('\n', ' '))
    md.append("| " + " | ".join(hdr_texts) + " |")
    
    # Separator
    separators = [":---" for _ in hdr_texts]
    md.append("| " + " | ".join(separators) + " |")
    
    # Data rows
    for row in table.rows[1:]:
        row_texts = []
        for cell in row.cells:
            cell_text = " ".join(get_markdown_text(p) for p in cell.paragraphs).strip()
            row_texts.append(cell_text.replace('\n', ' '))
        md.append("| " + " | ".join(row_texts) + " |")
        
    return "\n".join(md)

def convert_docx_to_md(docx_path: Path, md_path: Path) -> None:
    if not docx_path.exists():
        print(f"Error: {docx_path} does not exist.")
        return
        
    doc = Document(docx_path)
    md_content = []
    
    # Add metadata block or similar if needed, else start writing content
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            style_name = item.style.name if item.style else "Normal"
            text = get_markdown_text(item).strip()
            
            if style_name == "Title":
                md_content.append(f"# {text}\n")
            elif style_name == "Heading 1":
                md_content.append(f"## {text}\n")
            elif style_name == "Heading 2":
                md_content.append(f"### {text}\n")
            elif style_name == "Heading 3":
                md_content.append(f"#### {text}\n")
            elif style_name == "Heading 4":
                md_content.append(f"##### {text}\n")
            elif style_name == "List Bullet":
                md_content.append(f"* {text}")
            else:
                # Normal paragraph or empty line
                if not text:
                    md_content.append("")
                else:
                    md_content.append(f"{text}\n")
        elif isinstance(item, Table):
            table_md = format_table_as_md(item)
            md_content.append(table_md + "\n")
            
    # Write to file
    with open(md_path, "w", encoding="utf-8") as f:
        # Join with double newlines for paragraphs, but single for lists/tables where appropriate
        # To make it simple, write each block separated by appropriate spacing
        output_str = ""
        for block in md_content:
            if block.startswith("* "):
                output_str += block + "\n"
            elif block.startswith("|"):
                output_str += block + "\n"
            else:
                output_str += block + "\n"
        f.write(output_str)
        
    print(f"Converted docx -> md successfully: {md_path}")

if __name__ == "__main__":
    convert_docx_to_md(DOCX_PATH, MD_OUTPUT_PATH)
