import os
import re
import base64
import zlib
import urllib.request
import ssl
import uuid
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def download_mermaid_image(mermaid_code):
    try:
        # Clean up lines and check if empty
        lines = [line.rstrip() for line in mermaid_code.splitlines() if line.strip()]
        if not lines:
            return None
        cleaned_code = "\n".join(lines)
        
        # Kroki requires standard deflate compression + URL-safe base64
        compressed = zlib.compress(cleaned_code.encode('utf-8'))
        b64 = base64.urlsafe_b64encode(compressed).decode('utf-8')
        url = f"https://kroki.io/mermaid/png/{b64}"
        
        # Bypass local SSL check issues if any
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        )
        
        with urllib.request.urlopen(req, context=ctx, timeout=15) as response:
            img_data = response.read()
            temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "mermaid_images")
            os.makedirs(temp_dir, exist_ok=True)
            filename = f"mermaid_{uuid.uuid4().hex[:8]}.png"
            filepath = os.path.join(temp_dir, filename)
            with open(filepath, "wb") as f:
                f.write(img_data)
            return filepath
    except Exception as e:
        print(f"Error downloading Mermaid image: {e}")
        return None

def add_code_fallback(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code.rstrip())
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(60, 60, 60)

def add_formatted_text(paragraph, text):
    # Parse inline formatting: bold (**text**), code (`text`), link ([text](url))
    pattern = r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 50, 50)
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 90, 180)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_to_doc(doc, table_data):
    if not table_data:
        return
    
    # Check if first row is header
    has_header = len(table_data) > 1 and all(cell.strip() == '-' or cell.strip().startswith(':') or cell.strip().endswith(':') or all(c == '-' for c in cell.strip()) for cell in table_data[1])
    
    rows_to_add = []
    if has_header:
        headers = table_data[0]
        data_rows = table_data[2:]
        rows_to_add.append((headers, True))
        for r in data_rows:
            rows_to_add.append((r, False))
    else:
        for r in table_data:
            rows_to_add.append((r, False))
            
    num_cols = len(rows_to_add[0][0])
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'
    
    for row_idx, (row_cells, is_header) in enumerate(rows_to_add):
        row = table.add_row()
        for col_idx, cell_text in enumerate(row_cells):
            if col_idx >= len(row.cells):
                break
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            clean_text = cell_text.strip()
            add_formatted_text(p, clean_text)
            
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if is_header:
                set_cell_background(cell, "0055A5")  # Dark Blue header
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
            elif row_idx % 2 == 0:
                set_cell_background(cell, "F2F5F9")  # Zebra striping light blue/gray
                
def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return
        
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 Width
    section.page_height = Inches(11.69) # A4 Height
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Style setups
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(30, 30, 30)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Table detection
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_lines.append(cells)
            i += 1
            continue
        elif in_table:
            add_table_to_doc(doc, table_lines)
            table_lines = []
            in_table = False
            
        # Empty line
        if not stripped:
            i += 1
            continue
            
        # Heading 1
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 85, 165) # Premium Dark Blue
            
        # Heading 2
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 204) # Lighter Blue
            
        # Heading 3
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(80, 80, 80) # Dark Gray
            
        # Heading 4
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(11.5)
            
        # Blockquote
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            
        # List items (Bullet list)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = leading_spaces // 2
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            content = stripped[2:]
            add_formatted_text(p, content)
            
        # Numbered list items
        elif re.match(r'^\d+\.\s+', stripped):
            match = re.match(r'^(\d+)\.\s+(.*)', stripped)
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = leading_spaces // 2
            
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            content = match.group(2)
            add_formatted_text(p, content)
            
        # Mermaid code blocks - Render and download image using Kroki API (supports unicode/accents)
        elif stripped.startswith('```mermaid'):
            i += 1
            mermaid_code = ""
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_code += lines[i]
                i += 1
            
            if mermaid_code:
                img_path = download_mermaid_image(mermaid_code)
                if img_path and os.path.exists(img_path):
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(6)
                        
                        # Add image to word doc
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(5.8))
                        
                        # Add caption
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_after = Pt(12)
                        caption_run = caption_p.add_run("Sơ đồ trực quan hóa (Hệ thống Trạm Cân)")
                        caption_run.italic = True
                        caption_run.font.size = Pt(9.0)
                        caption_run.font.color.rgb = RGBColor(120, 120, 120)
                    except Exception as img_err:
                        print(f"Error inserting Mermaid image: {img_err}")
                        add_code_fallback(doc, mermaid_code)
                else:
                    # Fallback to plain code formatting
                    add_code_fallback(doc, mermaid_code)
            
        # Code block opening/closing
        elif stripped.startswith('```'):
            i += 1
            code_text = ""
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_text += lines[i]
                i += 1
            if code_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text.rstrip())
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(60, 60, 60)
            
        # Horizontal rule
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("•  •  •")
            
        # Image detection (e.g., ![Màn hình Đăng nhập](images/login.png))
        elif stripped.startswith('![') and ']' in stripped and stripped.endswith(')'):
            match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
            if match:
                caption = match.group(1)
                img_rel_path = match.group(2)
                # Resolve path relative to md_path
                md_dir = os.path.dirname(os.path.abspath(md_path))
                abs_img_path = os.path.normpath(os.path.join(md_dir, img_rel_path))
                if os.path.exists(abs_img_path):
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(6)
                        
                        # Add image
                        run = p.add_run()
                        run.add_picture(abs_img_path, width=Inches(5.8))
                        
                        # Add caption
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_after = Pt(12)
                        caption_run = caption_p.add_run(caption)
                        caption_run.italic = True
                        caption_run.font.size = Pt(9.5)
                        caption_run.font.color.rgb = RGBColor(120, 120, 120)
                    except Exception as img_err:
                        print(f"Error inserting local image {abs_img_path}: {img_err}")
                        p = doc.add_paragraph()
                        p.add_run(stripped)
                else:
                    print(f"Warning: Image file not found: {abs_img_path}")
                    p = doc.add_paragraph()
                    p.add_run(stripped)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_text(p, stripped)
                
        # Normal Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, stripped)
            
        i += 1
        
    # Flush remaining table if file ends during table parsing
    if in_table and table_lines:
        add_table_to_doc(doc, table_lines)
        
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == '__main__':
    base_dir = r'g:\Source-code\pmcan_C#\SRSdocs'
    
    # Convert BRD
    brd_md = os.path.join(base_dir, 'StationApp_System_BRD.md')
    brd_docx = os.path.join(base_dir, 'StationApp_System_BRD.docx')
    convert_md_to_docx(brd_md, brd_docx)
    
    # Convert SRS
    srs_md = os.path.join(base_dir, 'StationApp_System_SRS.md')
    srs_docx = os.path.join(base_dir, 'StationApp_System_SRS.docx')
    convert_md_to_docx(srs_md, srs_docx)
    
    print("All Word document generation completed!")
