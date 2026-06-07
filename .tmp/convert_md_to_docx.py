import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def add_formatted_text(paragraph, text):
    # Parse inline formatting: bold (**text**), code (`text`), link ([text](url))
    # Simplify regex parsing
    pattern = r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            # Add a subtle gray highlight if possible (handled simply by font color in python-docx)
            run.font.color.rgb = RGBColor(180, 50, 50)
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 90, 180)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def set_cell_background(cell, fill_hex):
    # Set background color of a cell
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # Set margins (padding) of a cell in twentieths of a point (dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_to_doc(doc, table_data):
    # table_data is a list of lists of strings
    if not table_data:
        return
    
    # Check if first row is header
    has_header = len(table_data) > 1 and all(cell.strip() == '-' or cell.strip().startswith(':') or cell.strip().endswith(':') or all(c == '-' for c in cell.strip()) for cell in table_data[1])
    
    rows_to_add = []
    if has_header:
        headers = table_data[0]
        # skip table_data[1] because it is the divider |---|---|
        data_rows = table_data[2:]
        rows_to_add.append((headers, True))
        for r in data_rows:
            rows_to_add.append((r, False))
    else:
        for r in table_data:
            rows_to_add.append((r, False))
            
    num_cols = len(rows_to_add[0][0])
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'
    
    for row_idx, (row_cells, is_header) in enumerate(rows_to_add):
        row = table.add_row()
        for col_idx, cell_text in enumerate(row_cells):
            cell = row.cells[col_idx]
            cell.text = "" # Clear default text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            clean_text = cell_text.strip()
            add_formatted_text(p, clean_text)
            
            # Formatting cells
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if is_header:
                set_cell_background(cell, "0055A5")  # Dark Blue header
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
            elif row_idx % 2 == 0:
                set_cell_background(cell, "F2F5F9")  # Zebra striping light blue/gray
                
def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return
        
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 Width
    section.page_height = Inches(11.69) # A4 Height
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Style setups
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(30, 30, 30)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    in_list = False
    list_level = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Table detection
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            # Parse cells
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_lines.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            add_table_to_doc(doc, table_lines)
            table_lines = []
            in_table = False
            # Fall through to process current line
            
        # Empty line
        if not stripped:
            i += 1
            continue
            
        # Heading 1
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 85, 165) # Premium Dark Blue
            
        # Heading 2
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 204) # Lighter Blue
            
        # Heading 3
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(80, 80, 80) # Dark Gray
            
        # Heading 4
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(11.5)
            
        # Blockquote
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add left border style simulated by background highlight or bold/italic
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            
        # List items (Bullet list)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            # Calculate indent based on leading spaces of raw line
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = leading_spaces // 2
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            content = stripped[2:]
            add_formatted_text(p, content)
            
        # Numbered list items
        elif re.match(r'^\d+\.\s+', stripped):
            match = re.match(r'^(\d+)\.\s+(.*)', stripped)
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = leading_spaces // 2
            
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            content = match.group(2)
            add_formatted_text(p, content)
            
        # Mermaid code blocks (ignore visual markup in word doc)
        elif stripped.startswith('```mermaid'):
            # Skip until closing block
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            
        # Code block opening/closing
        elif stripped.startswith('```'):
            # Skip code blocks or add as monospaced paragraph
            i += 1
            code_text = ""
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_text += lines[i]
                i += 1
            if code_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text.rstrip())
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(60, 60, 60)
            
        # Horizontal rule
        elif stripped == '---':
            # Add a subtle line or page break? Let's just do space
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("•  •  •")
            
        # Normal Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, stripped)
            
        i += 1
        
    # Flush remaining table if file ends during table parsing
    if in_table and table_lines:
        add_table_to_doc(doc, table_lines)
        
    # Cover page or title block styling (let's insert title at the top of document)
    # python-docx doesn't easily support prepending paragraphs cleanly, but we can do it:
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == '__main__':
    md_file = r'g:\Source-code\pmcan_C#\SRSdocs\StationApp_System_BRD.md'
    docx_file = r'g:\Source-code\pmcan_C#\SRSdocs\StationApp_System_BRD.docx'
    convert_md_to_docx(md_file, docx_file)
